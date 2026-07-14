r"""
Real-template accuracy suite for LetterCheck Smart.

The suite uses real templates from D:\Study\Diplom\data as the source corpus.
For each suitable template it creates mutated copies with one intentional
mistake and checks that the analyzer finds that mistake. If the mistake is not
found, the assertion fails and the printed accuracy goes down.

Run from backend:
    python tests/accuracy_suite.py --template-dir D:\Study\Diplom\data
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_LINE_SPACING

try:
    from openpyxl import load_workbook
except ImportError:  # XLSX reference material is optional for the accuracy run.
    load_workbook = None

os.environ.setdefault("LETTERCHECK_DISABLE_LLM", "1")

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from extract_docx_header import extract_header_and_greeting  # noqa: E402


GENERATED_DIR = Path(__file__).resolve().parent / "generated_realistic_cases"
REPORT_PATH = Path(__file__).resolve().parent / "accuracy_report.json"


@dataclass
class AssertionSpec:
    category: str
    name: str
    check: Callable[[dict[str, Any]], tuple[bool, Any]]
    expected: Any


@dataclass
class AccuracyCase:
    name: str
    source: Path
    target: Path
    mutation: str
    assertions: list[AssertionSpec] = field(default_factory=list)
    note: str = ""


def clean(text: str | None) -> str:
    return " ".join((text or "").replace("\xa0", " ").split())


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen:
                    continue
                seen.add(marker)
                for paragraph in cell.paragraphs:
                    yield paragraph


def clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.add_run("")


def replace_first_text(doc: Document, pattern: str, replacement: str, flags: int = re.IGNORECASE) -> bool:
    regex = re.compile(pattern, flags)
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        if regex.search(text):
            paragraph.text = regex.sub(replacement, text, count=1)
            return True
    return False


def clear_first_text(doc: Document, predicate: Callable[[str], bool]) -> bool:
    for paragraph in iter_all_paragraphs(doc):
        if predicate(clean(paragraph.text)):
            clear_paragraph(paragraph)
            return True
    return False


def collect_template_paths(template_dir: Path) -> list[Path]:
    return [
        path for path in sorted(template_dir.glob("*.docx"))
        if "выдержка" not in path.name.lower()
    ]


def read_reference_material(template_dir: Path) -> dict[str, Any]:
    instruction = next(template_dir.glob("*Инструкции*.docx"), None)
    requirements = next(template_dir.glob("*.xlsx"), None)
    keywords = ("адрес", "обращ", "исполн", "подпис", "прилож", "интервал", "исход")

    instruction_hits = []
    if instruction:
        doc = Document(str(instruction))
        for paragraph in doc.paragraphs:
            text = clean(paragraph.text)
            if text and any(keyword in text.lower() for keyword in keywords):
                instruction_hits.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(clean(cell.text) for cell in row.cells if clean(cell.text))
                if row_text and any(keyword in row_text.lower() for keyword in keywords):
                    instruction_hits.append(row_text)

    requirement_rows = []
    requirements_note = None
    if requirements and load_workbook:
        workbook = load_workbook(requirements, data_only=True)
        for ws in workbook.worksheets:
            for row in ws.iter_rows(values_only=True):
                values = [clean(str(value)) if value is not None else "" for value in row]
                if any(values):
                    requirement_rows.append(values[:8])
    elif requirements:
        requirements_note = "openpyxl не установлен, XLSX-требования не прочитаны"

    return {
        "instruction_file": instruction.name if instruction else None,
        "instruction_hits": instruction_hits[:80],
        "requirements_file": requirements.name if requirements else None,
        "requirements_rows": requirement_rows[:80],
        "requirements_note": requirements_note,
    }


def make_case_name(index: int, source: Path, mutation: str) -> str:
    return f"{index:03d}_{source.stem}__{mutation}.docx"


def save_mutation(source: Path, target: Path, mutator: Callable[[Document, dict[str, Any]], bool], baseline: dict[str, Any]) -> bool:
    doc = Document(str(source))
    changed = mutator(doc, baseline)
    if not changed:
        return False
    doc.save(str(target))
    return True


def status_starts(path: str, prefix: str) -> AssertionSpec:
    def check(result: dict[str, Any]) -> tuple[bool, Any]:
        current: Any = result
        for part in path.split("."):
            current = current.get(part, {}) if isinstance(current, dict) else {}
        return isinstance(current, str) and current.startswith(prefix), current

    return AssertionSpec(path.split(".")[0], path, check, prefix)


def error_block(block_name: str) -> AssertionSpec:
    def check(result: dict[str, Any]) -> tuple[bool, Any]:
        errors = result.get("check_required_blocks", {}).get("errors", [])
        blocks = [error.get("block") for error in errors]
        return block_name in blocks, blocks

    return AssertionSpec("required_blocks", f"missing_{block_name}", check, block_name)


def error_type(check_name: str, expected_type: str) -> AssertionSpec:
    def check(result: dict[str, Any]) -> tuple[bool, Any]:
        errors = result.get(check_name, {}).get("errors", [])
        types = [error.get("type") for error in errors]
        return expected_type in types, types

    return AssertionSpec(check_name, f"{check_name}_{expected_type}", check, expected_type)


def status_equals(path: str, expected: str) -> AssertionSpec:
    def check(result: dict[str, Any]) -> tuple[bool, Any]:
        current: Any = result
        for part in path.split("."):
            current = current.get(part, {}) if isinstance(current, dict) else {}
        return current == expected, current

    return AssertionSpec(path.split(".")[0], path, check, expected)


def has_greeting_and_person(baseline: dict[str, Any]) -> bool:
    return bool(baseline.get("greeting") and baseline.get("parsed_greeting", {}).get("persons"))


def mutate_wrong_gender(doc: Document, baseline: dict[str, Any]) -> bool:
    greeting = baseline.get("greeting") or ""
    if "Уважаемая" in greeting:
        return replace_first_text(doc, r"\bУважаемая\b", "Уважаемый")
    if "Уважаемый" in greeting:
        return replace_first_text(doc, r"\bУважаемый\b", "Уважаемая")
    if "Уважаемые" in greeting:
        return replace_first_text(doc, r"\bУважаемые\b", "Уважаемый")
    return False


def mutate_wrong_initial(doc: Document, baseline: dict[str, Any]) -> bool:
    persons = baseline.get("parsed_greeting", {}).get("persons") or []
    if not persons:
        return False
    name = persons[0].get("name")
    if not name:
        return False
    replacement = "Петр" if not name.startswith("П") else "Иван"
    return replace_first_text(doc, rf"\b{re.escape(name)}\b", replacement)


def mutate_remove_greeting(doc: Document, baseline: dict[str, Any]) -> bool:
    return clear_first_text(doc, lambda text: text.lower().startswith("уважа"))


def mutate_remove_addressee(doc: Document, baseline: dict[str, Any]) -> bool:
    if doc.tables and len(doc.tables[0].rows) > 0 and len(doc.tables[0].rows[0].cells) > 1:
        if clean(doc.tables[0].cell(0, 1).text):
            doc.tables[0].cell(0, 1).text = ""
            return True

    addressee = clean(baseline.get("addressee") or "")
    if not addressee:
        return False
    fio = baseline.get("fio") or []
    markers = []
    markers.extend(person.get("surname") for person in fio if person.get("surname"))
    markers.extend(line for line in addressee.splitlines() if len(clean(line)) > 8)
    markers = [clean(marker) for marker in markers if clean(marker)]

    def is_addressee(text: str) -> bool:
        lower = text.lower()
        if lower.startswith("уважа") or "федеральная служба" in lower or "rosreestr" in lower:
            return False
        return any(marker in text for marker in markers)

    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen:
                    continue
                seen.add(marker)
                if is_addressee(clean(cell.text)):
                    cell.text = ""
                    return True
    return clear_first_text(doc, is_addressee)


def mutate_remove_signature(doc: Document, baseline: dict[str, Any]) -> bool:
    raw_baseline_texts = [
        item.get("text") or ""
        for item in baseline.get("blocks", {}).get("signature_paragraphs", [])
        if clean(item.get("text"))
    ]
    baseline_texts = [clean(text) for text in raw_baseline_texts]
    if baseline_texts:
        baseline_parts = {
            clean(part)
            for text in raw_baseline_texts
            for part in re.split(r"\s*\n\s*|\s{2,}", text)
            if clean(part)
        }
        changed = False
        for paragraph in iter_all_paragraphs(doc):
            text = clean(paragraph.text)
            if text in baseline_texts or text in baseline_parts:
                clear_paragraph(paragraph)
                changed = True
        return changed

    signature_words = ("руководител", "начальник", "заместител", "директор", "представител")
    fio_re = re.compile(r"\b[А-ЯЁ]\.[ \t]*[А-ЯЁ]\.?\s+[А-ЯЁ][а-яё]+")
    changed = False
    for paragraph in iter_all_paragraphs(doc):
        text = clean(paragraph.text)
        lower = text.lower()
        if any(word in lower for word in signature_words) or (len(text) <= 90 and fio_re.search(text)):
            clear_paragraph(paragraph)
            changed = True
    return changed


def mutate_remove_executor(doc: Document, baseline: dict[str, Any]) -> bool:
    baseline_texts = [
        clean(item.get("text"))
        for item in baseline.get("blocks", {}).get("executor_paragraphs", [])
        if clean(item.get("text"))
    ]
    if baseline_texts:
        changed = False
        for paragraph in iter_all_paragraphs(doc):
            text = clean(paragraph.text)
            if text in baseline_texts:
                clear_paragraph(paragraph)
                changed = True
        return changed

    changed = False
    for paragraph in iter_all_paragraphs(doc):
        text = clean(paragraph.text)
        if (
            "исполнитель" in text.lower()
            or bool(re.search(r"\b8\s*\(\d{3}\)", text))
            or bool(re.search(r"\b\d{2,4}-\d{2}-\d{2}\b", text))
        ):
            clear_paragraph(paragraph)
            changed = True
    return changed


def mutate_remove_attachment(doc: Document, baseline: dict[str, Any]) -> bool:
    regex = re.compile(r"\bПриложени[ея]:", re.IGNORECASE)
    changed = False
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        if regex.search(text):
            paragraph.text = regex.sub("Материалы:", text)
            changed = True
    return changed


def mutate_bad_normative_number(doc: Document, baseline: dict[str, Any]) -> bool:
    bad_sentence = (
        " В соответствии с Федеральным законом от 13.07.2015 "
        "«О государственной регистрации недвижимости» просим направить сведения."
    )
    greeting = clean(baseline.get("greeting"))
    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)
        if text and text != greeting and len(text) > 60:
            paragraph.text = paragraph.text + bad_sentence
            return True
    return False


def mutate_bad_line_spacing(doc: Document, baseline: dict[str, Any]) -> bool:
    greeting = baseline.get("greeting")
    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)
        if text and text != clean(greeting) and len(text) > 40:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = 3.0
            return True
    return False


def add_case(
    cases: list[AccuracyCase],
    source: Path,
    mutation: str,
    mutator: Callable[[Document, dict[str, Any]], bool],
    baseline: dict[str, Any],
    assertions: list[AssertionSpec],
    note: str = "",
) -> None:
    target = GENERATED_DIR / make_case_name(len(cases) + 1, source, mutation)
    if save_mutation(source, target, mutator, baseline):
        cases.append(AccuracyCase(
            name=target.stem,
            source=source,
            target=target,
            mutation=mutation,
            assertions=assertions,
            note=note,
        ))


def build_cases(template_dir: Path) -> list[AccuracyCase]:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    for old_file in GENERATED_DIR.glob("*.docx"):
        old_file.unlink()

    cases: list[AccuracyCase] = []
    for source in collect_template_paths(template_dir):
        baseline = extract_header_and_greeting(str(source))
        required_status = baseline.get("check_required_blocks", {}).get("status")
        layout_status = baseline.get("check_layout", {}).get("status")

        baseline_target = GENERATED_DIR / make_case_name(len(cases) + 1, source, "original")
        shutil.copy2(source, baseline_target)
        cases.append(AccuracyCase(
            name=baseline_target.stem,
            source=source,
            target=baseline_target,
            mutation="original",
            assertions=[
                status_equals("check_required_blocks.status", "ok"),
                status_equals("check_layout.status", "ok"),
            ],
            note="Оригинальный шаблон из корпуса не должен давать ошибок обязательных блоков и интервала.",
        ))

        if has_greeting_and_person(baseline):
            add_case(
                cases, source, "wrong_gender", mutate_wrong_gender, baseline,
                [status_starts("check_gender.status", "error_gender")],
                "Меняем род/число слова обращения.",
            )
            add_case(
                cases, source, "wrong_initial", mutate_wrong_initial, baseline,
                [status_starts("check_initials.status", "error_mismatch")],
                "Меняем имя в обращении так, чтобы первая буква не совпала с инициалом адресата.",
            )

        if baseline.get("check_required_blocks", {}).get("document_type") == "outgoing_letter" and baseline.get("greeting"):
            add_case(
                cases, source, "remove_greeting", mutate_remove_greeting, baseline,
                [error_block("has_greeting"), status_starts("check_gender.status", "skip")],
                "Удаляем обращение из письма, где оно обязательно.",
            )

        if baseline.get("addressee"):
            add_case(
                cases, source, "remove_addressee", mutate_remove_addressee, baseline,
                [error_block("has_addressee")],
                "Удаляем адресата из шапки.",
            )

        if required_status == "ok" and baseline.get("blocks", {}).get("detected", {}).get("has_signature"):
            add_case(
                cases, source, "remove_signature", mutate_remove_signature, baseline,
                [error_block("has_signature")],
                "Удаляем реквизит подписи.",
            )

        if required_status == "ok" and baseline.get("blocks", {}).get("detected", {}).get("has_executor"):
            add_case(
                cases, source, "remove_executor", mutate_remove_executor, baseline,
                [error_block("has_executor")],
                "Удаляем отметку об исполнителе/телефон.",
            )

        full_text = (baseline.get("blocks", {}).get("full_text") or "").lower()
        if "приложение" in source.name.lower() and "приложение:" in full_text:
            add_case(
                cases, source, "remove_attachment", mutate_remove_attachment, baseline,
                [error_block("has_attachment")],
                "Переименовываем реквизит «Приложение:», чтобы блок перестал распознаваться.",
            )

        if baseline.get("check_spelling", {}).get("status") == "ok" and "№" in (baseline.get("blocks", {}).get("full_text") or ""):
            add_case(
                cases, source, "bad_normative_number", mutate_bad_normative_number, baseline,
                [status_starts("check_spelling.status", "has_errors"), error_type("check_spelling", "normative")],
                "Удаляем номер из нормативной ссылки.",
            )

        if layout_status == "ok":
            add_case(
                cases, source, "bad_line_spacing", mutate_bad_line_spacing, baseline,
                [status_starts("check_layout.status", "has_errors"), error_type("check_layout", "line_spacing")],
                "Ставим в основном тексте тройной межстрочный интервал.",
            )

    return cases


def run_case(case: AccuracyCase) -> dict[str, Any]:
    result = extract_header_and_greeting(str(case.target))
    assertion_results = []
    for spec in case.assertions:
        passed, actual = spec.check(result)
        assertion_results.append({
            "category": spec.category,
            "name": spec.name,
            "passed": passed,
            "expected": spec.expected,
            "actual": actual,
        })
    return {
        "name": case.name,
        "source": str(case.source),
        "file": str(case.target),
        "mutation": case.mutation,
        "note": case.note,
        "assertions": assertion_results,
        "observed": {
            "check_gender": result.get("check_gender", {}).get("status"),
            "check_initials": result.get("check_initials", {}).get("status"),
            "check_spelling": result.get("check_spelling", {}).get("status"),
            "check_required_blocks": result.get("check_required_blocks", {}).get("status"),
            "check_layout": result.get("check_layout", {}).get("status"),
            "document_type": result.get("check_required_blocks", {}).get("document_type"),
        },
    }


def print_case_line(report: dict[str, Any]) -> None:
    failed = [item for item in report["assertions"] if not item["passed"]]
    status = "PASS" if not failed else "FAIL"
    total = len(report["assertions"])
    passed = total - len(failed)
    print(f"{status:4} {report['mutation'][:22]:22} {passed:2}/{total:2}  {Path(report['source']).name}")
    for item in failed:
        print(
            f"     - [{item['category']}] {item['name']}: "
            f"expected={item['expected']!r}, actual={item['actual']!r}"
        )


def summarize(assertions: list[dict[str, Any]]) -> dict[str, Any]:
    total = len(assertions)
    passed = sum(1 for item in assertions if item["passed"])
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for item in assertions:
        grouped[item["category"]].append(item)

    by_category = {}
    for category, items in sorted(grouped.items()):
        cat_total = len(items)
        cat_passed = sum(1 for item in items if item["passed"])
        by_category[category] = {
            "passed": cat_passed,
            "total": cat_total,
            "accuracy": round(cat_passed / cat_total, 4) if cat_total else 0,
        }

    return {
        "passed": passed,
        "total": total,
        "accuracy": round(passed / total, 4) if total else 0,
        "by_category": by_category,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Mutate real DOCX templates and measure analyzer accuracy.")
    parser.add_argument("--template-dir", type=Path, default=Path(r"D:\Study\Diplom\data"))
    args = parser.parse_args()

    cases = build_cases(args.template_dir)
    all_assertions: list[dict[str, Any]] = []
    case_reports = []

    print("\nReal-template mutation accuracy cases")
    print("-" * 96)
    for case in cases:
        report = run_case(case)
        print_case_line(report)
        case_reports.append(report)
        all_assertions.extend(report["assertions"])

    summary = summarize(all_assertions)
    print("-" * 96)
    print(f"TOTAL accuracy: {summary['passed']}/{summary['total']} = {summary['accuracy']:.2%}")
    for category, data in summary["by_category"].items():
        print(f"{category:22}: {data['passed']}/{data['total']} = {data['accuracy']:.2%}")

    full_report = {
        "summary": summary,
        "template_dir": str(args.template_dir),
        "generated_dir": str(GENERATED_DIR),
        "llm_disabled": os.getenv("LETTERCHECK_DISABLE_LLM") == "1",
        "reference_material": read_reference_material(args.template_dir),
        "cases": case_reports,
    }
    REPORT_PATH.write_text(json.dumps(full_report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nReport saved: {REPORT_PATH}")

    return 0 if summary["passed"] == summary["total"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
