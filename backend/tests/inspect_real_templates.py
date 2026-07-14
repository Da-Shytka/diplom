from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


def clean(text: str) -> str:
    return " ".join((text or "").replace("\xa0", " ").split())


def docx_summary(path: Path) -> dict:
    doc = Document(str(path))
    tables = []
    for table in doc.tables[:3]:
        rows = []
        for row in table.rows[:8]:
            cells = []
            seen = set()
            for cell in row.cells:
                value = clean(cell.text)
                if value and value not in seen:
                    seen.add(value)
                    cells.append(value)
            if cells:
                rows.append(cells)
        tables.append(rows)

    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    greeting = next((p for p in paragraphs if p.lower().startswith("уважа")), None)
    if not greeting:
        for table_rows in tables:
            for row in table_rows:
                for cell in row:
                    if cell.lower().startswith("уважа"):
                        greeting = cell
                        break

    return {
        "file": path.name,
        "tables": tables,
        "paragraphs_head": paragraphs[:25],
        "paragraphs_tail": paragraphs[-12:],
        "greeting": greeting,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
    }


def xlsx_summary(path: Path) -> dict:
    wb = load_workbook(path, data_only=False)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 80), values_only=True):
            values = [clean(str(v)) if v is not None else "" for v in row]
            if any(values):
                rows.append(values[:12])
        sheets.append({"sheet": ws.title, "rows": rows[:60]})
    return {"file": path.name, "sheets": sheets}


def main() -> int:
    source = Path(sys.argv[1] if len(sys.argv) > 1 else r"D:\Study\Diplom\data")
    out = Path(sys.argv[2] if len(sys.argv) > 2 else Path(__file__).with_name("template_inventory.json"))
    docs = []
    sheets = []
    for path in sorted(source.iterdir()):
        if path.suffix.lower() == ".docx":
            docs.append(docx_summary(path))
        elif path.suffix.lower() == ".xlsx":
            sheets.append(xlsx_summary(path))

    out.write_text(json.dumps({"source": str(source), "docx": docs, "xlsx": sheets}, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Inventory saved: {out}")
    print(f"DOCX: {len(docs)}, XLSX: {len(sheets)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
